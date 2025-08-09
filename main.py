import os
from tempfile import NamedTemporaryFile
from flask import Flask, send_file, jsonify
from google.cloud import firestore, storage
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

# Initialize Firestore and Storage clients
db = firestore.Client(database='dawood')
storage_client = storage.Client()
BUCKET_NAME = 'internship-2025-465209.appspot.com'

def create_project_page(doc, project):
    """Adds a project to the document as a new page"""
    # Add project title
    title = doc.add_heading(project['title'], level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add description
    doc.add_paragraph("Description:", style='Heading 2')
    doc.add_paragraph(project.get('description', 'N/A'))
    
    # Add supervisor info
    doc.add_paragraph("Supervisor:", style='Heading 2')
    doc.add_paragraph(project.get('supervisorId', 'N/A'))
    
    # Add students
    if 'students' in project:
        doc.add_paragraph("Team Members:", style='Heading 2')
        for student in project['students']:
            doc.add_paragraph(f"• {student['name']} ({student['email']})")
    
    # Add page break (except after last project)
    doc.add_page_break()

@app.route('/generate-projects-report', methods=['GET'])
def generate_report():
    """Endpoint to generate and return DOCX report"""
    try:
        print("Starting report generation...")
        
        # 1. Get all projects from Firestore
        projects_ref = db.collection('projects')
        docs = projects_ref.stream()
        
        # 2. Create new document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # 3. Add each project to the document
        for doc_num, doc_snapshot in enumerate(docs):
            project = doc_snapshot.to_dict()
            print(f"Adding project {doc_num+1}: {project.get('title')}")
            create_project_page(doc, project)
        
        # 4. Save to temporary file
        with NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            print(f"Temporary file created at: {tmp.name}")
            
            # 5. Upload to Google Storage (optional)
            blob = storage_client.bucket(BUCKET_NAME).blob('reports/all_projects.docx')
            blob.upload_from_filename(tmp.name)
            print(f"Uploaded to Storage: {blob.public_url}")
            
            # 6. Return the file for download
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name='all_projects.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        return jsonify({"error": str(e)}), 500
        
    finally:
        # Clean up temporary file
        if 'tmp' in locals():
            os.unlink(tmp.name)
            print("Temporary file removed")

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)), debug=True)